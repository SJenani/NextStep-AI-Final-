import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_from_data(data: dict) -> bytes:
    """
    Generates a professional DOCX resume from structured data.
    """
    doc = Document()
    
    # Configure document settings
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Standardize font to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
        
    # 1. Name
    name = data.get('name', 'Candidate Name')
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(16)
    
    # 2. Contact Info
    contact_parts = []
    if data.get('email'): contact_parts.append(data['email'])
    if data.get('linkedin'): contact_parts.append(data['linkedin'])
    if data.get('github'): contact_parts.append(data['github'])
    
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_contact = p_contact.add_run(" | ".join(contact_parts))
        r_contact.font.size = Pt(10)
        
    # 3. Role
    role = data.get('role')
    if role:
        p_role = doc.add_paragraph()
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_role = p_role.add_run(role)
        r_role.italic = True
        r_role.font.size = Pt(12)
        
    # Helper to add section headings
    def add_heading(text: str):
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
        
    # 4. Summary
    summary = data.get('summary')
    if summary:
        add_heading('Summary')
        doc.add_paragraph(summary)
        
    # 5. Experience
    experience = data.get('experience', [])
    if experience:
        add_heading('Experience Highlights')
        for exp in experience:
            # Experience might be string lines or complex objects.
            # Assuming strings based on structured_profile in schemas
            if isinstance(exp, str):
                doc.add_paragraph(exp, style='List Bullet')
            else:
                doc.add_paragraph(str(exp), style='List Bullet')
            
    # 6. Projects
    projects = data.get('projects', [])
    if projects:
        add_heading('Projects')
        for proj in projects:
            if isinstance(proj, str):
                doc.add_paragraph(proj, style='List Bullet')
            else:
                doc.add_paragraph(str(proj), style='List Bullet')
            
    # 7. Skills & Tools
    skills = data.get('skills', [])
    tools = data.get('tools', [])
    all_skills = skills + tools
    if all_skills:
        add_heading('Technical Skills')
        doc.add_paragraph(", ".join(all_skills))
        
    # 8. Certifications
    certs = data.get('certifications', [])
    if certs:
        add_heading('Certifications & Education')
        for cert in certs:
            if isinstance(cert, str):
                doc.add_paragraph(cert, style='List Bullet')
            else:
                doc.add_paragraph(str(cert), style='List Bullet')
            
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
